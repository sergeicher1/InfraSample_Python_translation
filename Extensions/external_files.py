# ------------------------------------------------------------------------------------------------
# -- coding                                   | utf-8
# -- Author                                   | Sergei Chernyahovsky
# -- Site                                     | http://sergeicher.pro/
# -- Favorite Quote                           | “Always code as if the guy who ends up
#                                                   maintaining your code will be a violent
#                                                       psychopath who knows where you live”
# -- Language                                 | Python
# -- Version                                  | 3.11
# -- WebDriver                                | Selenium
# -- Version                                  | 4.6.0
# -- Description                              | External files
# Usage Example
# Import this module/class and use these functions
# ------------------------------------------------------------------------------------------------


import allure
import xml.etree.ElementTree as ET
import PyPDF2
import openpyxl
import json
import csv
import docx
from docx import Document
from openpyxl.drawing.image import Image
from openpyxl import load_workbook


class XML:

    # Parameter   : Name of the attribute in xml file
    # Return value: String
    @staticmethod
    @allure.step("Read data from XML file")
    def ReadData(dataPath: str, name: str) -> str:
        root = ET.parse(dataPath).getroot()
        return root.find(".//" + name).text


############################################################################

# continue from here implementing as needed


############################################################################


class PDF:

    @staticmethod
    @allure.step("Read PDF file")
    def Read(filePath: str):
        pdf_file_obj = open(filePath, 'rb')
        pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)
        # printing number of pages in pdf file
        print("Number of pages: ", pdf_reader.numPages)
        page_obj = pdf_reader.getPage(0)
        text_from_pdf = page_obj.extractText()
        pdf_file_obj.close()
        return text_from_pdf


############################################################################

# continue from here implementing as needed


############################################################################


class Text:

    @staticmethod
    @allure.step("Adds text to txt file, creates new if file doesn't exist")
    def Add(filePath: str, data: str):
        with open(filePath, "a") as file:
            file.writelines(data)
            file.write("\n")

    '''Read text file, File must exist'''

    @staticmethod
    @allure.step("Reads text file")
    def Read(filePath: str):
        with open(filePath, "r") as file:
            text = file.read()
        return text


############################################################################

# continue from here implementing as needed


############################################################################


class Excel:

    # Create workbook object
    # Get workbook active sheet object from the active attribute Cell objects also
    # have a row, column, and coordinate attributes that provide
    # location information for the cell.
    @staticmethod
    @allure.step("Load workbook object")
    def LoadWorkbookObject(filePath: str):
        try:
            excel_obj = openpyxl.load_workbook(filePath)
            sheet_obj = excel_obj.active
            return sheet_obj
        except FileExistsError:
            print(f"File {filePath} doesn't exists")

    '''
    Reading from Spreadsheets, Multiple Cells
    Give the location of the file to open the workbook, workbook object is created
    Get workbook active sheet object from the active attribute Cell objects
    We can get the count of the total rows and columns using the
    max_row and max_column respectively. We can use these values inside the
    for loop to get the value of the desired row or column or any cell depending upon
    the situation. Let’s see how to get the value of the first column and first row.
    '''

    @staticmethod
    @allure.step("Read column from excel file")
    def ReadColumn(filePath: str, col: int) -> list:
        row = Excel.LoadWorkbookObject(filePath=filePath).max_row
        result = []
        for i in range(1, row + 1):
            cell_obj = Excel.LoadWorkbookObject(filePath=filePath).cell(row=i, column=col)
            result.append(cell_obj.value)
        return result[1:]

    @staticmethod
    @allure.step("Read row from excel file")
    def ReadRow(filePath: str, row: int) -> list:
        column = Excel.LoadWorkbookObject(filePath=filePath).max_column
        result = []
        for i in range(1, column + 1):
            cell_obj = Excel.LoadWorkbookObject(filePath=filePath).cell(row=row, column=i)
            result.append(cell_obj.value)
        return result

    @staticmethod
    @allure.step("Read active sheet from excel file")
    def ReadSpreadsheets(filePath: str):
        cell_obj = Excel.LoadWorkbookObject(filePath=filePath)
        result = []
        for i in cell_obj:
            for j in i:
                result.append(j.value)
        return result

    @staticmethod
    @allure.step("Read value from specific cell in excel file")
    def ReadCellsValue(filePath: str, rowInt: int, colInt: int) -> str:
        cell_obj = Excel.LoadWorkbookObject(filePath=filePath).cell(row=rowInt, column=colInt)
        return cell_obj.value

    '''
    Writing to Spreadsheets
    '''

    # REWRITE !!! or Creates a new file
    @staticmethod
    @allure.step("Saves value to excel file, if file exists -> rewrites")
    def SaveValueNewXL(filePath: str, val: str):
        excel_obj = openpyxl.Workbook()
        sheet_obj = excel_obj.active
        val1 = sheet_obj.cell(row=1, column=1)
        val1.value = val
        excel_obj.save(filePath)

    # File must be created first
    # CELL Example : path, "A1" or "B6", "ABCD"
    @staticmethod
    @allure.step("Update value in specific cell in excel file")
    def UpdateValueInCell(filePath: str, cell: str, value: str):
        try:
            excel_obj = load_workbook(filename=filePath)
            sheet_obj = excel_obj.active
            sheet_obj[cell] = value
            excel_obj.save(filename=filePath)
        except FileExistsError:
            print(f"File {filePath} doesn't exists")

        # finally: TODO: check if needed
        #     print(print(f"File {filePath} doesn't exist, \n Message: "))
        #         # print(f"File {filePath} doesn't exist, \n Message: ", m)

    '''Adding Images
    For the purpose of importing images inside our worksheet,
     we would be using openpyxl.drawing.image.Image. The method is a wrapper over
     PIL.Image method found in PIL (pillow) library. Due to which it is necessary
     for the PIL (pillow) library to be installed in order to use this method.
    Image Used:'''

    @staticmethod
    @allure.step("Saves image in new excel file")
    def SaveImageInNewXL(filePath: str, imagePath: str, cell: str):
        wb = openpyxl.Workbook()
        sheet = wb.active
        img = Image(imagePath)
        sheet.add_image(img, cell)
        wb.save(filePath)

    @staticmethod
    @allure.step("Saves image in existing excel file")
    def SaveImage(filePath: str, imagePath: str, cell: str):
        try:
            wb = openpyxl.load_workbook(filePath)
            sheet = wb.active
            img = Image(imagePath)
            sheet.add_image(img, cell)
            wb.save(filePath)
        except FileExistsError:
            print(f"File {filePath} doesn't exists")


############################################################################

# continue from here implementing as needed


############################################################################


class JSON:

    @staticmethod
    @allure.step("Create JSON file")
    def Create(filePath: str, value: dict):
        with open(filePath, "w") as p:
            json.dump(value, p)

    '''Read existing Json file'''

    @staticmethod
    @allure.step("Read JSON file")
    def Read(filePath: str) -> dict:
        with open(filePath, "r") as reader:
            data = json.load(reader)
        return data

    '''Update values in Json file'''

    # The file must exist
    # before this function specify data values to be changed
    # Example :
    # data = JsonRead(path)
    # data["Salary"]["QA"] = "10000"
    # data["Salary"]["Dev"] = "10000"
    @staticmethod
    @allure.step("Change values in JSON file")
    def ChangeValue(filePath: str, dat):
        JSON.Create(filePath, dat)


############################################################################

# continue from here implementing as needed


############################################################################


class Docx:

    @staticmethod
    @allure.step("Create and write to DOCX file")
    def CreateAndAdd(filePath: str, text: str):
        doc = docx.Document()
        doc_para = doc.add_paragraph()
        doc_para.add_run(text)
        doc.add_page_break()
        doc.save(filePath)

    '''Read from docx file'''

    @staticmethod
    @allure.step("Read the DOCX file")
    def Read(filePath: str):
        doc = Document(filePath)
        result = []
        for para in doc.paragraphs:
            result.append(para.text)
        return result


############################################################################

# continue from here implementing as needed


############################################################################


class CSV:

    @staticmethod
    @allure.step("CSV ReadAll")
    def ReadAll(filePath: str):
        result = []
        with open(filePath, mode='r') as file:
            csv_file = csv.reader(file)
            for lines in csv_file:
                result.insert(len(result), lines)
            return result

    '''Write to CSV file'''

    @staticmethod
    @allure.step("CSV Write to file")
    def WriteTo(filePath: str, fields: list, rows: list):
        with open(filePath, 'w') as csvfile:
            csv_writer = csv.writer(csvfile)
            csv_writer.writerow(fields)
            csv_writer.writerows(rows)

############################################################################

# continue from here implementing as needed

############################################################################
