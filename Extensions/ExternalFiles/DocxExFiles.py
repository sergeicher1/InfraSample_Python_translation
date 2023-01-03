# ------------------------------------------------------------------------------------------------
# -- coding                                   | utf-8
# -- Author                                   | Sergei Chernyahovsky
# -- Site                                     | http://sergeicher.pro/
# -- Favorite Quote                           | “Always code as if the guy who ends up
#                                                   maintaining your code will be a violent
#                                                       psychopath who knows where you live”
# -- Language                                 | Python
# -- Version                                  | 3.11
# -- WebDriver                                | Selenium
# -- Version                                  | 4.6.0
# -- Description                              | Docx External files
# Usage Example
# Import this module and use these functions
# ------------------------------------------------------------------------------------------------

import allure
import docx
from docx import Document

class Docx:

    @staticmethod
    @allure.step("Create and write to DOCX file")
    def CreateAndAdd(filePath: str, text: str):
        doc = docx.Document()
        doc_para = doc.add_paragraph()
        doc_para.add_run(text)
        doc.add_page_break()
        doc.save(filePath)


    '''Read from docx file'''

    @staticmethod
    @allure.step("Read the DOCX file")
    def Read(filePath: str):
        doc = Document(filePath)
        result = []
        for para in doc.paragraphs:
            result.append(para.text)
        return result

############################################################################

# continue from here implementing as needed


############################################################################
